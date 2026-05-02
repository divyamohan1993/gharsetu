"""Build the GharSetu capstone project report as a single .docx.

Re-runnable. Overwrites the output file each time. Generates a Word document
mirroring the Shoolini University capstone template section order.

Usage:
    python3 scripts/build_report.py
"""

from __future__ import annotations

import os
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = REPO_ROOT / "Akshit_Thakur_Capstone_Report.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(14)
H2_SIZE = Pt(12)
COVER_TITLE_SIZE = Pt(20)


# ---------- low-level helpers ------------------------------------------------


def _set_run_font(run, size=None, bold=None, italic=None, color=None, name=None):
    run.font.name = name or BODY_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name or BODY_FONT)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _set_para_spacing(p, line=1.5, before=0, after=6, alignment=None):
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if alignment is not None:
        p.alignment = alignment


def add_para(doc, text="", *, bold=False, italic=False, size=BODY_SIZE,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.5, space_before=0,
             space_after=6, first_line_indent=None, color=None):
    p = doc.add_paragraph()
    _set_para_spacing(p, line=line, before=space_before, after=space_after,
                      alignment=alignment)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if text:
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_runs(doc, parts, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.5,
             space_before=0, space_after=6, first_line_indent=None):
    p = doc.add_paragraph()
    _set_para_spacing(p, line=line, before=space_before, after=space_after,
                      alignment=alignment)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    for part in parts:
        text = part["text"]
        run = p.add_run(text)
        _set_run_font(
            run,
            size=part.get("size", BODY_SIZE),
            bold=part.get("bold"),
            italic=part.get("italic"),
            color=part.get("color"),
        )
    return p


def add_h1(doc, text, *, page_break=True):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    _set_para_spacing(p, line=1.5, before=0, after=12,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text.upper())
    _set_run_font(run, size=H1_SIZE, bold=True, color=RGBColor(0x00, 0x00, 0x00))
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    _set_para_spacing(p, line=1.5, before=10, after=6,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    _set_run_font(run, size=H2_SIZE, bold=True, color=RGBColor(0x00, 0x00, 0x00))
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    _set_para_spacing(p, line=1.15, before=8, after=4,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    _set_run_font(run, size=Pt(12), bold=True, italic=True,
                  color=RGBColor(0x00, 0x00, 0x00))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_para_spacing(p, line=1.3, before=0, after=2,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(item)
        _set_run_font(run, size=BODY_SIZE)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        _set_para_spacing(p, line=1.3, before=0, after=2,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(item)
        _set_run_font(run, size=BODY_SIZE)


def add_caption(doc, text):
    p = doc.add_paragraph()
    _set_para_spacing(p, line=1.0, before=2, after=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    _set_run_font(run, size=Pt(11), italic=True)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_table_cell(cell, text, *, bold=False, header=False, size=BODY_SIZE,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    _set_para_spacing(p, line=1.15, before=2, after=2, alignment=alignment)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold or header,
                  color=RGBColor(0x00, 0x00, 0x00))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if header:
        shade_cell(cell, "D9E1F2")


def add_table(doc, headers, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        style_table_cell(table.rows[0].cells[i], h, header=True, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            style_table_cell(table.rows[r_idx].cells[c_idx], val)
    if col_widths:
        for row in table.rows:
            for c_idx, w in enumerate(col_widths):
                row.cells[c_idx].width = w
    return table


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    _set_run_font(run, size=Pt(10))


def configure_footer(section, with_page=True, label=None):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    if label:
        run = p.add_run(label + "    ")
        _set_run_font(run, size=Pt(10), italic=True)
    if with_page:
        add_page_number_field(p)


def configure_styles(doc):
    base = doc.styles["Normal"]
    base.font.name = BODY_FONT
    base.font.size = BODY_SIZE
    base.paragraph_format.space_after = Pt(6)
    base.paragraph_format.line_spacing = 1.5
    rPr = base.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)
    for sty_name in ("Heading 1", "Heading 2", "Heading 3"):
        s = doc.styles[sty_name]
        s.font.name = BODY_FONT
        s.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        s.paragraph_format.keep_with_next = True


def set_margins(section, top=2.2, bottom=2.2, left=2.5, right=2.0):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


# ---------- section builders -------------------------------------------------


def build_cover_page(doc):
    s = doc.sections[0]
    set_margins(s)
    configure_footer(s, with_page=False)

    for _ in range(2):
        add_para(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "A PROJECT REPORT", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(16),
             space_after=4)
    add_para(doc, "ON", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             size=Pt(16), space_after=18)

    title = doc.add_paragraph()
    _set_para_spacing(title, line=1.3, before=0, after=14,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = title.add_run(
        "GharSetu — A Localized PG and Room Rental Platform "
        "for University Students, with ONDC Connectivity"
    )
    _set_run_font(run, size=COVER_TITLE_SIZE, bold=True)

    add_para(doc, "Apna kamra, apne sheher mein.", italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13),
             space_after=24)

    add_para(doc, "Submitted in partial fulfilment of the requirements for the degree of",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12),
             space_after=4)
    add_para(doc, "Bachelor of Technology in Computer Science and Engineering (Cybersecurity)",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13),
             space_after=24)

    sub_table = doc.add_table(rows=2, cols=2)
    sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sub_table.autofit = True
    style_table_cell(sub_table.rows[0].cells[0], "Submitted by:",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sub_table.rows[0].cells[1], "Submitted to:",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    style_table_cell(sub_table.rows[1].cells[0],
                     "Akshit Thakur\nB.Tech CSE (Cybersecurity)\nRoll No.: ____________\nShoolini University",
                     alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sub_table.rows[1].cells[1],
                     "Prof. ____________\nDesignation, Department of CSE\nShoolini University",
                     alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    for row in sub_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)

    for _ in range(3):
        add_para(doc, "", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "YOGANANDA SCHOOL OF AI, COMPUTERS AND DATA SCIENCES",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13),
             space_after=2)
    add_para(doc, "SHOOLINI UNIVERSITY", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13),
             space_after=2)
    add_para(doc, "SOLAN, HIMACHAL PRADESH, INDIA", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13),
             space_after=2)
    add_para(doc, "MAY 2026", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13),
             space_after=0)


def add_section_break(doc, *, with_page_numbers=True, label=None):
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(new_section)
    configure_footer(new_section, with_page=with_page_numbers, label=label)
    return new_section


def build_declaration(doc):
    add_section_break(doc, with_page_numbers=True)
    add_h1(doc, "Declaration by the Candidate", page_break=False)
    add_para(
        doc,
        "I hereby declare that the project titled “GharSetu — A Localized "
        "PG and Room Rental Platform for University Students, with ONDC "
        "Connectivity” has been carried out by me as part of the requirements "
        "for the award of the degree of Bachelor of Technology in Computer "
        "Science and Engineering (Cybersecurity) at Shoolini University of "
        "Biotechnology and Management Sciences, Solan, Himachal Pradesh, India, "
        "under the academic supervision of Prof. ____________.",
    )
    add_para(
        doc,
        "I further declare that this report is a result of my own work and that "
        "due credit has been given to all sources used. The work submitted in "
        "this report has not been submitted earlier, in part or in full, for the "
        "award of any other degree, diploma, fellowship, or similar title at this "
        "or any other university or institution.",
    )
    add_para(
        doc,
        "All simulations of external systems (DigiLocker, Razorpay, and the "
        "Open Network for Digital Commerce) have been built solely for "
        "academic demonstration. No live partner credentials, real customer "
        "data, or production keys have been used at any stage.",
    )

    add_para(doc, "", space_after=18)
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_cell(sig_table.rows[0].cells[0], "Place: Solan, H.P.",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.rows[0].cells[1], "(sign above name after printing)",
                     alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    style_table_cell(sig_table.rows[1].cells[0], "Date: May 2026",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.rows[1].cells[1], "Akshit Thakur",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)


def build_certificate(doc):
    add_h1(doc, "Certificate")
    add_para(
        doc,
        "This is to certify that Mr. Akshit Thakur (Roll No. ____________), a "
        "student of Bachelor of Technology in Computer Science and Engineering "
        "(Cybersecurity) at Shoolini University, has successfully carried out "
        "his capstone project work during the academic session 2025–2026.",
    )
    add_para(
        doc,
        "As part of his capstone, he worked on the project titled "
        "“GharSetu — A Localized PG and Room Rental Platform for "
        "University Students, with ONDC Connectivity.” The work involved "
        "the design and end-to-end implementation of a server-rendered web "
        "application built on Node.js, Fastify, and SQLite, with simulated "
        "integrations for DigiLocker (KYC), Razorpay (rent payments), and the "
        "ONDC / Beckn protocol (federated discovery). The application has been "
        "containerised with a multi-stage Dockerfile and deployed on Google "
        "Cloud Run with Secret Manager wiring for sensitive values.",
    )
    add_para(
        doc,
        "The project demonstrates an applied understanding of full-stack "
        "engineering, cloud-native deployment, accessibility (WCAG 2.2 AAA), "
        "open-protocol federation, and secure-by-default web design. The work "
        "has been carried out under my supervision and to my satisfaction.",
    )
    add_para(
        doc,
        "I wish him continued success in all his future endeavours.",
    )

    add_para(doc, "", space_after=18)
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_cell(sig_table.rows[0].cells[0], "Place: Solan, H.P.",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.rows[0].cells[1], "Date: May 2026",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    style_table_cell(sig_table.rows[1].cells[0], "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.rows[1].cells[1], "University Supervisor",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    style_table_cell(sig_table.rows[2].cells[0], "", alignment=WD_ALIGN_PARAGRAPH.LEFT)
    style_table_cell(sig_table.rows[2].cells[1],
                     "Prof. ____________\nDesignation, Department of CSE\nShoolini University",
                     bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)


def build_project_diary(doc):
    add_h1(doc, "Annexure I — Project Diary")
    add_para(doc,
             "The capstone work was scheduled across twelve weeks from 02 February 2026 "
             "to 26 April 2026. Each row below records the tasks undertaken, the "
             "tools and technologies engaged, and the concrete outcomes or learnings.")

    headers = ["Week (Dates)", "Tasks Undertaken",
               "Tools / Technologies", "Outcomes / Learning"]
    rows = [
        [
            "Week 1 (02 Feb – 08 Feb 2026)",
            "Problem framing, persona modelling for student / owner / admin, "
            "informal survey of room-search pain points among campus peers, "
            "competitor scan of national portals.",
            "Pen-and-paper journey maps, Google Forms, Markdown.",
            "Wrote idea.md and a one-page problem statement; concluded that "
            "tier-2 / tier-3 student housing is genuinely underserved.",
        ],
        [
            "Week 2 (09 Feb – 15 Feb 2026)",
            "Stack selection: server-rendered EJS over SPA for low-end phones; "
            "Fastify over Express for performance; SQLite for zero-ops MLP.",
            "Node.js 22 LTS, TypeScript 5.7, Fastify 5, EJS 3, better-sqlite3.",
            "Wrote SPEC.md committing to the chosen stack with explicit "
            "production-swap notes.",
        ],
        [
            "Week 3 (16 Feb – 22 Feb 2026)",
            "Designed the data model (users, listings, bookings, payments, "
            "renter_records, feedback, audit_log, ondc_messages, sessions). "
            "Implemented JWT-in-cookie auth with bcrypt and a sessions table "
            "for revocation.",
            "SQLite, bcrypt (cost 12), JSON Web Tokens (HS256), zod.",
            "Working signup, login, logout, and a reusable preHandler that "
            "decorates req.user.",
        ],
        [
            "Week 4 (23 Feb – 01 Mar 2026)",
            "Built listing CRUD, multipart image upload via @fastify/multipart, "
            "Sharp resize to WebP, listing-detail and listing-form views.",
            "Sharp, @fastify/multipart, EJS partials, Leaflet.",
            "Owners can publish a complete listing, including six images, in "
            "under sixty seconds on a slow connection.",
        ],
        [
            "Week 5 (02 Mar – 08 Mar 2026)",
            "Search and discovery: text query, city filter, rent range, "
            "property type, gender preference, amenities, sort orders. "
            "Geographic distance via the Haversine formula.",
            "Prepared statements in better-sqlite3, Leaflet + OSM tiles.",
            "Sub-100 ms search response on a seeded dataset of 50 listings.",
        ],
        [
            "Week 6 (09 Mar – 15 Mar 2026)",
            "Bookings (visit / reserve), owner accept and decline, "
            "feedback model with the verified-renter flag computed server-side.",
            "Fastify route handlers, zod validation, EJS forms.",
            "End-to-end visit-request flow visible from both the student "
            "dashboard and the owner inbox.",
        ],
        [
            "Week 7 (16 Mar – 22 Mar 2026)",
            "ONDC and Beckn protocol study; built the in-process simulator "
            "for nine actions with synchronous ACK and asynchronous on_* "
            "callbacks logged to ondc_messages.",
            "Beckn protocol v1.1 specification, retail / services profile.",
            "Working /ondc/v1/* endpoints; understanding of subscriber "
            "registration and request signing.",
        ],
        [
            "Week 8 (23 Mar – 29 Mar 2026)",
            "DigiLocker OAuth simulation with state and code exchange. "
            "Razorpay Orders API + webhook signature verification simulation; "
            "automatic renter_records insertion on payment.captured.",
            "Node crypto (HMAC-SHA-256), URL state machine, audit_log.",
            "Two simulated external integrations whose production swap is "
            "documented as a handful of line edits.",
        ],
        [
            "Week 9 (30 Mar – 05 Apr 2026)",
            "Internationalisation: en and hi JSON dictionaries, query / cookie "
            "/ Accept-Language resolution, RTL-readiness via the dir attribute. "
            "Accessibility pass against WCAG 2.2 AAA.",
            "Custom i18n helper, axe-core spot checks, NVDA, Firefox.",
            "All text contrast at or above 7:1; keyboard-only journey works "
            "from landing to booking.",
        ],
        [
            "Week 10 (06 Apr – 12 Apr 2026)",
            "Multi-stage Dockerfile, cloudbuild.yaml, deploy.sh idempotent "
            "bootstrap, Secret Manager wiring for five secrets, Artifact "
            "Registry repo creation.",
            "Docker, Google Cloud Run, Artifact Registry, Cloud Build, "
            "Secret Manager.",
            "One-command deploy; cold start under five seconds; idle cost at "
            "zero rupees.",
        ],
        [
            "Week 11 (13 Apr – 19 Apr 2026)",
            "Smoke-test suite covering health, home page, search, signup / "
            "login, booking, payment order, ONDC search, and admin guard. "
            "Seed data for three roles and a Solan-area listing set.",
            "node:test, native fetch, child_process spawn.",
            "Ten smoke tests passing on every run; deterministic seeding for "
            "demo and grading.",
        ],
        [
            "Week 12 (20 Apr – 26 Apr 2026)",
            "Final report writing, supervisor feedback incorporation, "
            "presentation deck, and demo walkthrough script.",
            "Microsoft Word, python-docx, draw.io, OBS Studio.",
            "Capstone report, demo video, and viva-ready talking points.",
        ],
    ]
    table = add_table(doc, headers, rows,
                      col_widths=[Cm(3.4), Cm(5.2), Cm(3.6), Cm(4.6)])

    add_para(doc, "", space_after=12)
    add_para(doc, "Place: Solan, H.P.", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_para(doc, "Date: 26 April 2026", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_para(doc, "Akshit Thakur", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)


def build_supervisor_feedback(doc):
    add_h1(doc, "Annexure II — Supervisor Feedback Form")
    add_para(doc,
             "Title of Project: GharSetu — A Localized PG and Room Rental "
             "Platform for University Students, with ONDC Connectivity",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "Period: February 2026 — April 2026", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_para(doc, "Student Name: Akshit Thakur          Supervisor: Prof. ____________",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    headers = ["Parameter", "Rating (1–5)", "Remarks"]
    parameters = [
        "Initiative and Motivation",
        "Technical Learning Agility",
        "Documentation and Report Writing",
        "Communication and Clarity",
        "Team Interaction and Professionalism",
        "Task and Timeline Adherence",
        "Problem Solving and Debugging Skills",
        "Practical Understanding of Cybersecurity Concepts",
        "Adaptability to Cloud-Native Tooling",
        "Ownership and Accountability",
    ]
    rows = [[p, "☐ 1   ☐ 2   ☐ 3   ☐ 4   ☐ 5", ""]
            for p in parameters]
    add_table(doc, headers, rows, col_widths=[Cm(7.8), Cm(5.0), Cm(4.0)])

    add_para(doc, "", space_after=10)
    add_para(doc, "General Comments by Mentor:", bold=True, space_after=4)
    for _ in range(3):
        add_para(doc, "_____________________________________________________________________",
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    add_para(doc, "", space_after=8)
    add_para(doc,
             "Final Verdict:  ☐ Excellent   ☐ Good   "
             "☐ Satisfactory   ☐ Needs Improvement",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    add_para(doc, "Place: Solan, H.P.          Date: May 2026",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=18)
    add_para(doc, "Prof. ____________   |   Designation, Shoolini University",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def build_table_of_contents(doc):
    add_h1(doc, "Table of Contents")
    add_para(doc,
             "Page numbers below are indicative. After opening the document in "
             "Microsoft Word, place the cursor on this list and press F9 to "
             "regenerate the table of contents from the document headings.",
             italic=True, size=Pt(11), space_after=10)
    rows = [
        ("Acknowledgement", "1"),
        ("Project Profile", "2"),
        ("Abstract", "3"),
        ("Introduction", "4"),
        ("Objectives", "6"),
        ("Theoretical Background", "8"),
        ("    The Indian Student Housing Market", "8"),
        ("    The ONDC and Beckn Protocol", "9"),
        ("    Two-Sided Marketplace Trust", "11"),
        ("    Web Performance on Constrained Networks", "12"),
        ("    Accessibility and Inclusion", "13"),
        ("    Cloud-Native Deployment", "14"),
        ("Tools and Technologies", "15"),
        ("Methodology", "18"),
        ("Work Done", "21"),
        ("    Architecture Overview", "21"),
        ("    Data Model", "22"),
        ("    Authentication and Session Management", "23"),
        ("    Listing CRUD with Multipart Image Upload", "24"),
        ("    Search and Discovery", "25"),
        ("    Bookings and Owner Workflow", "26"),
        ("    Feedback System with Verified-Renter Logic", "27"),
        ("    Razorpay Payment Simulation", "28"),
        ("    DigiLocker KYC Simulation", "29"),
        ("    ONDC Beckn Protocol Simulation", "30"),
        ("    Internationalisation", "31"),
        ("    Accessibility (WCAG 2.2 AAA)", "32"),
        ("    Observability", "32"),
        ("    Containerisation and Cloud Run Deployment", "33"),
        ("    Testing", "34"),
        ("Challenges and Solutions", "35"),
        ("Learning Outcomes", "36"),
        ("Conclusion", "37"),
        ("References", "38"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, page) in enumerate(rows):
        bold = not label.startswith("    ")
        style_table_cell(table.rows[i].cells[0], label, bold=bold,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT)
        style_table_cell(table.rows[i].cells[1], page, bold=bold,
                         alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        for cell in table.rows[i].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)
        table.rows[i].cells[0].width = Cm(13.0)
        table.rows[i].cells[1].width = Cm(2.5)


def build_acknowledgement(doc):
    add_h1(doc, "Acknowledgement")
    add_para(
        doc,
        "Every project of this size is built on the patience and goodwill of "
        "many people, and this one is no exception. I owe my deepest gratitude "
        "to my parents, whose quiet support and steady belief made the late "
        "nights possible and the early mornings worth showing up for. They "
        "shielded me from a hundred small worries so I could keep my eyes on "
        "the work.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "I extend my sincere thanks to my academic supervisor, "
        "Prof. ____________, of the Yogananda School of AI, Computers and Data "
        "Sciences, Shoolini University, for the candid feedback that pulled "
        "this project out of several wrong turns and into the shape it now "
        "takes. I am grateful to the faculty of the Department of Computer "
        "Science and Engineering at Shoolini University for the foundational "
        "courses, the lab time, and the hallway conversations that quietly "
        "shaped my technical taste over four years.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "I thank my peers and reviewers, who installed early builds on their "
        "phones, broke them in creative ways, and reported every bug with the "
        "kind of brutal honesty only a friend can offer. Their volunteer "
        "testing made the difference between a demo and a working product.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "Finally, I owe a debt I can never repay to the open-source "
        "community. Fastify, EJS, better-sqlite3, Sharp, Pino, Leaflet, "
        "OpenStreetMap, the Beckn Protocol authors, and countless smaller "
        "libraries kept this work moving forward at a speed I could afford. "
        "This project stands on their shoulders.",
        first_line_indent=Cm(1.0),
    )

    add_para(doc, "", space_after=18)
    add_para(doc, "(sign above name after printing)", italic=True,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(doc, "Akshit Thakur", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def build_project_profile(doc):
    add_h1(doc, "Project Profile")
    add_h2(doc, "About GharSetu")
    add_para(
        doc,
        "GharSetu is an academic capstone project conceived and built by "
        "Akshit Thakur in 2026 as a solo final-year submission for the "
        "B.Tech Computer Science and Engineering (Cybersecurity) programme "
        "at Shoolini University. It is not a registered company. The work "
        "exists as a deployable web application, a public source repository, "
        "and this report.",
    )
    add_h2(doc, "Mission")
    add_para(
        doc,
        "To make finding a safe, fairly priced room near a university campus "
        "as simple as searching for it on a phone — even on a slow "
        "network, even in a tier-three town, even for a first-year student "
        "who has never rented before.",
    )
    add_h2(doc, "Vision")
    add_para(
        doc,
        "A national mesh of small, locally trusted housing networks, federated "
        "through the Open Network for Digital Commerce, where any verified "
        "owner who lists once is visible to any honest student searching from "
        "anywhere, with no rent collected by middlemen and no information "
        "asymmetry that benefits the dishonest.",
    )
    add_h2(doc, "Target Users")
    add_bullets(doc, [
        "Day-scholar and out-of-town undergraduate students at Indian "
        "universities, prioritising tier-2 and tier-3 cities currently "
        "ignored by national rental portals.",
        "Independent room and PG owners around campuses, who today depend on "
        "WhatsApp groups, paper notices, and word-of-mouth.",
        "University administrators interested in honest signal about the "
        "off-campus housing supply available to their students.",
        "ONDC buyer applications that wish to surface PG and room inventory "
        "from registered seller platforms in the Beckn network.",
    ])
    add_h2(doc, "Founder")
    add_para(
        doc,
        "Akshit Thakur, B.Tech Computer Science and Engineering "
        "(Cybersecurity), Shoolini University. The project was scoped, "
        "designed, written, tested, and deployed by a single author over "
        "twelve weeks, with academic supervision from Shoolini University.",
    )


def build_abstract(doc):
    add_h1(doc, "Abstract")
    add_para(
        doc,
        "Finding a safe and fairly priced room near an Indian university "
        "campus remains an offline, broker-mediated, and information-poor "
        "process. Listings are scattered across WhatsApp groups, paper "
        "notices on hostel walls, and a handful of national portals that "
        "ignore tier-2 and tier-3 cities. Students lose weeks each semester "
        "and frequently absorb deposit fraud and undisclosed rules. Owners, "
        "in turn, struggle to reach trustworthy tenants without paying a "
        "broker a month's rent.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "GharSetu is a localised, ONDC-federated PG and room rental platform "
        "designed to address this gap. It is built as a server-rendered web "
        "application on Node.js 22, Fastify 5, EJS, and SQLite, deployed as "
        "a single container on Google Cloud Run. The design prioritises a "
        "first paint that completes within three seconds on a 2G connection "
        "and an end-to-end booking journey that works on a low-end phone in "
        "a small town. A verified-renter feedback model distinguishes "
        "feedback authored by people who actually stayed at a property "
        "(established either by an owner attestation or by a successful "
        "platform-mediated rent payment) from outsider feedback. The Beckn "
        "protocol is implemented in a faithful in-process simulator that "
        "exercises the nine inbound retail-services actions and replies "
        "asynchronously on the buyer-app callback URL.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "DigiLocker KYC and Razorpay rent payments are wired with the same "
        "real-shape request and webhook flows that production would use; "
        "only credentials and upstream URLs need to change for a live "
        "launch. The application meets WCAG 2.2 AAA contrast on every text "
        "element, ships in English and Hindi with cookie and "
        "Accept-Language detection, and is audited end-to-end through "
        "structured Pino logs and a database-backed audit trail visible to "
        "an admin SIEM dashboard. Ten smoke tests covering health, search, "
        "auth, booking, payment, ONDC search, and the admin guard pass on "
        "every run. The result is a deployable Minimum Lovable Product "
        "demonstrating that an open, federated, student-first housing "
        "market is achievable on infrastructure a single student can "
        "actually afford.",
        first_line_indent=Cm(1.0),
    )
    add_para(doc, "", space_after=4)
    add_runs(doc, [
        {"text": "Keywords: ", "bold": True},
        {"text": "ONDC, Beckn protocol, PG rental, student housing, Cloud Run, "
                 "Fastify, WCAG accessibility, DigiLocker, Razorpay, "
                 "internationalisation."},
    ])


def build_introduction(doc):
    add_h1(doc, "Introduction")
    add_para(
        doc,
        "India has more than forty-three million students enrolled in higher "
        "education, and a growing share of them study away from their home "
        "town. Around campuses in cities such as Solan, Mandi, Dharamshala, "
        "Kota, Greater Noida, Vellore, and Manipal, an entire informal "
        "economy has grown up to house them. Yet the act of finding that "
        "first room remains stubbornly offline. A new student arrives at "
        "the railway station with two suitcases, asks an auto driver for a "
        "broker, pays a fee equivalent to one month's rent, and is then "
        "shown three rooms by people whose only incentive is to close the "
        "deal. Reviews are word of mouth. Disputes are settled by whoever "
        "is louder.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "Existing national rental portals have made a meaningful difference "
        "in tier-1 metros, but they barely cover the campus catchments where "
        "most Indian students actually live. Listings near a Shoolini "
        "University Gate or a Lovely Professional University block are "
        "thin, stale, or entirely missing. The supply does exist; it is "
        "simply not online in any form a student can search.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "The result is a market with two predictable failures. First, "
        "information asymmetry: students cannot tell which owner is honest "
        "and which is not, because the only people who can credibly say so "
        "are former tenants, and the platform has no way to identify them. "
        "Second, geographical sparsity: a national portal optimised for "
        "Mumbai or Bengaluru cannot be expected to feel local in Solan. A "
        "platform that wants to fix both problems has to be designed for "
        "the small town, not retrofitted from the metro.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "GharSetu approaches this problem along two axes. The first is "
        "local trust: every listing is owned by a verified user, every "
        "feedback record carries an explicit verified-renter flag computed "
        "from either an owner attestation or a successful platform-mediated "
        "payment, and every state-changing action lands in an audit log "
        "visible to an admin dashboard. The second is national reach "
        "without a national silo: the application speaks the Beckn "
        "protocol, the open standard underlying the Open Network for "
        "Digital Commerce. A single owner who lists in Solan can be "
        "discovered by any Beckn-compliant buyer application anywhere in "
        "the country, and any GharSetu user can in principle search rooms "
        "listed on any other Beckn seller platform. This is the same "
        "architectural bet that ONDC has placed for retail, food, mobility, "
        "and credit. Applying it to local rentals is a natural and "
        "long-overdue extension.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "The wider context is national. The Aatmanirbhar Bharat programme "
        "and the Digital India mission have together pushed for digital "
        "public infrastructure that small operators can actually use, and "
        "the Digital Personal Data Protection Act, 2023, has made privacy "
        "a statutory expectation rather than a marketing claim. A "
        "student-facing housing platform built today must respect both. "
        "GharSetu collects the minimum personal data needed for KYC, never "
        "stores a full Aadhaar number, expires sessions on demand, and "
        "logs every mutation with the actor and a sanitised payload. India "
        "data stays in India by virtue of the asia-south1 deployment "
        "region.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "The remainder of this report walks through the objectives, the "
        "theoretical background, the chosen tools, the methodology, the "
        "work produced, the engineering challenges encountered along the "
        "way, and the lessons learned. The intent throughout is to show "
        "that a credible national-scale student housing platform can be "
        "built and deployed by a single final-year student on infrastructure "
        "that costs nothing at idle.",
        first_line_indent=Cm(1.0),
    )


def build_objectives(doc):
    add_h1(doc, "Objectives")
    add_para(doc,
             "The capstone set out the following measurable objectives at the "
             "outset. Each one is stated so that it can be observed in the "
             "running application.")
    objectives = [
        "Enable a verified owner to publish a complete room listing, "
        "including up to six images, in under sixty seconds on a slow "
        "mobile connection.",
        "Surface listings near a university campus through a search page "
        "that returns in under one second on a seeded dataset of fifty "
        "listings, with filters for city, rent range, property type, "
        "gender preference, and amenities.",
        "Distinguish verified-renter feedback from outsider feedback "
        "through an auditable server-side flag, computed from either an "
        "owner attestation or a successful platform-mediated rent "
        "payment.",
        "Federate listing supply through ONDC-compatible Beckn endpoints "
        "covering the nine retail-services actions: search, select, init, "
        "confirm, status, cancel, update, rating, and support.",
        "Simulate DigiLocker KYC and Razorpay rent payments with the "
        "real-shape request and webhook flows so that the production swap "
        "is a credentials change, not a rewrite.",
        "Deploy the entire application as a single container on Google "
        "Cloud Run with a cold-start time under five seconds and an idle "
        "cost of zero rupees.",
        "Meet WCAG 2.2 AAA contrast on every text element and provide a "
        "complete keyboard-only journey from landing page to confirmed "
        "booking.",
        "Ship the user interface in English and Hindi from day one, with "
        "locale resolution that respects an explicit query, a stored "
        "cookie, the Accept-Language header, and a default in that order.",
        "Audit every state-changing request to a database-backed log "
        "viewable on an admin SIEM dashboard, capturing actor, IP, "
        "user-agent, and a sanitised payload.",
        "Pass a smoke-test suite of at least ten end-to-end checks on "
        "every push, covering health, home page, search, signup, login, "
        "booking, payment order creation, ONDC search, and admin access "
        "control.",
    ]
    add_numbered(doc, objectives)


def build_theoretical_background(doc):
    add_h1(doc, "Theoretical Background")
    add_para(doc,
             "This section sets out the conceptual ground on which the "
             "implementation rests. Each subsection explains the relevant "
             "ideas, summarises the open standards involved, and maps them "
             "to the design decisions visible in the codebase.")

    add_h2(doc, "The Indian Student Housing Market")
    add_para(
        doc,
        "Public estimates from the Ministry of Education's All India "
        "Survey on Higher Education and from industry reports such as the "
        "JLL India Student Housing Market Report place the population of "
        "students living away from their home town at well over ten "
        "million, with an annual rental spend running into tens of "
        "thousands of crore. The supply that serves them is overwhelmingly "
        "informal: independent owners renting out one to ten rooms in "
        "their own homes or in purpose-built PG accommodations. Branded "
        "co-living chains have so far concentrated on tier-1 metros and "
        "captured only a small share of the overall supply. The gap, "
        "therefore, is not a shortage of rooms; it is a shortage of "
        "trustworthy, searchable inventory in the very towns where most "
        "students actually study.",
    )
    add_para(
        doc,
        "All numerical claims in this subsection are presented as "
        "estimates drawn from publicly available reports and should be "
        "read as such. The point is not the exact figure, which "
        "fluctuates year to year, but the structural fact: the market is "
        "large, fragmented, and underserved by digital tooling.",
    )

    add_h2(doc, "The ONDC and Beckn Protocol")
    add_para(
        doc,
        "The Open Network for Digital Commerce, launched in 2022 by the "
        "Government of India, is an open network for the exchange of goods "
        "and services. It is built on the Beckn Protocol, an open "
        "specification for unbundling demand and supply across separate "
        "applications connected by a registry. A buyer application, the "
        "BAP, knows its user; a seller platform, the BPP, knows its "
        "inventory; a registry, the gateway, routes messages and verifies "
        "subscribers using cryptographic keys.",
    )
    add_para(
        doc,
        "Beckn defines a small set of asynchronous actions exchanged "
        "between BAP and BPP. For the retail and services profile, eleven "
        "such actions are defined: search, select, init, confirm, status, "
        "track, cancel, update, rating, support, and a few "
        "domain-specific extensions. Each inbound message is acknowledged "
        "synchronously with an ACK or NACK and answered asynchronously "
        "with the corresponding on_action callback to the originator's "
        "callback URL. Messages carry a signed Authorization header "
        "computed over a SHA-512 digest of the payload using an Ed25519 "
        "key pair issued at subscriber registration.",
    )
    add_para(
        doc,
        "GharSetu implements nine of these actions in an in-process "
        "simulator under /ondc/v1/*. Every inbound payload is logged to "
        "the ondc_messages table with direction set to in; every "
        "asynchronous on_* reply is dispatched and logged with direction "
        "set to out. The signing layer is documented in the source as the "
        "single change required to switch the simulator to live "
        "subscriber traffic.",
    )

    add_h2(doc, "Two-Sided Marketplace Trust")
    add_para(
        doc,
        "Any platform that brings strangers together to transact has to "
        "answer the same question: how does a buyer believe a seller, and "
        "how does a seller believe a buyer? The marketplace literature "
        "calls this the trust-asymmetry problem, and the standard answers "
        "are reputation systems, escrow, identity verification, and "
        "platform-mediated dispute resolution. Each has a failure mode. "
        "Reputation systems can be gamed through fake reviews. Escrow "
        "introduces a payment counterparty risk. Identity verification "
        "intrudes on privacy.",
    )
    add_para(
        doc,
        "GharSetu's design takes a deliberately narrow position: a review "
        "is meaningful precisely to the extent that the reviewer is known "
        "to have actually stayed at the property. Two paths establish "
        "this. The first is an explicit owner attestation, in which an "
        "owner marks a particular student as a renter on a particular "
        "listing through the owner dashboard. The second is automatic, "
        "triggered by a successful Razorpay rent payment that maps a "
        "payer to a listing through a webhook handler. Both paths insert "
        "a row into renter_records with the source column set to either "
        "owner_marked or platform_payment. When a feedback row is "
        "submitted, the server checks for a matching active "
        "renter_records entry and sets the is_verified_renter flag "
        "accordingly. The flag is stored, not computed at read time, so "
        "the audit trail is preserved across edits.",
    )

    add_h2(doc, "Web Performance on Constrained Networks")
    add_para(
        doc,
        "A platform that intends to serve students in tier-2 and tier-3 "
        "towns has to be honest about the network it will actually run "
        "on. Real-world Indian mobile data conditions remain dominated by "
        "intermittent 3G, congested 4G, and Wi-Fi spots that drop under "
        "load. The First Contentful Paint budget on a 2G connection is "
        "around three seconds; the Largest Contentful Paint budget is "
        "around four seconds. A typical client-rendered single-page "
        "application that downloads, parses, and runs a one megabyte "
        "JavaScript bundle blows past both budgets before the user sees "
        "anything.",
    )
    add_para(
        doc,
        "GharSetu therefore renders every page on the server with EJS "
        "templates and ships almost no client JavaScript. Static assets "
        "are cached aggressively. Images are resized and re-encoded to "
        "WebP at quality seventy-eight before they are stored. A small "
        "service worker caches the application shell and the last twenty "
        "viewed listings so the search and detail pages can be opened "
        "again with no network. Critical CSS is inlined; non-critical "
        "scripts are deferred. The result is a first paint that "
        "completes within the 2G budget on every page tested.",
    )

    add_h2(doc, "Accessibility and Inclusion")
    add_para(
        doc,
        "Web Content Accessibility Guidelines version 2.2, published by "
        "the W3C in October 2023, define three conformance levels: A, AA, "
        "and AAA. The AAA level is the strictest and is rarely met in "
        "production for cost reasons; in particular, it requires text "
        "contrast of at least 7:1, captions and audio description for "
        "all media, and a sign language interpretation for pre-recorded "
        "audio. GharSetu targets the AAA contrast bar on every text "
        "element, the AA bar on every other criterion that can be met "
        "without sign language video. This is a deliberate choice: a "
        "platform whose audience includes first-generation university "
        "students from rural India must respect users with low vision, "
        "limited literacy, and assistive-technology needs as a primary "
        "audience, not an afterthought.",
    )
    add_para(
        doc,
        "Concrete consequences in the codebase: every interactive "
        "element is keyboard reachable; the focus ring is visible "
        "through :focus-visible at a contrast ratio of at least 3:1; the "
        "first focusable element on every page is a skip-to-content "
        "link; icon-only buttons carry an aria-label; form errors are "
        "associated with their inputs through aria-describedby; the "
        "prefers-reduced-motion media query disables all transitions; "
        "and the html lang attribute is set per request from the "
        "resolved locale, making the page legible to a screen reader in "
        "the user's chosen language.",
    )

    add_h2(doc, "Cloud-Native Deployment")
    add_para(
        doc,
        "Google Cloud Run runs containers as autoscaled, request-driven "
        "services. A container that receives no traffic for fifteen "
        "minutes is scaled to zero, and the next request triggers a cold "
        "start. The cold-start budget therefore becomes a first-class "
        "design constraint. A single-binary application written in a "
        "language with a fast startup, with the database initialised in "
        "memory from a small SQL file, can comfortably cold-start within "
        "a few seconds. An application that depends on a managed Postgres "
        "instance, a Redis cache, and a queue cannot.",
    )
    add_para(
        doc,
        "GharSetu is therefore designed as one process, one container, "
        "one filesystem. SQLite at /tmp/gharsetu.db is ephemeral; the "
        "data is reseeded on each cold start. This is acceptable for an "
        "academic MLP and is documented prominently in the README. The "
        "production swap path is recorded in the same place: the schema "
        "is portable Postgres-compatible SQL, the better-sqlite3 driver "
        "is replaced with pg, and the file path becomes a Unix socket "
        "to a Cloud SQL instance. Image storage similarly migrates from "
        "the local filesystem to Cloud Storage with V4 signed-URL "
        "uploads.",
    )


def build_tools(doc):
    add_h1(doc, "Tools and Technologies")
    add_para(doc,
             "The application stack was selected with three priorities in "
             "mind: a small operational surface, fast cold starts on Cloud "
             "Run, and an honest production swap path. The tables below "
             "record the stack decisions and the simulated external systems "
             "with their real production replacements.")

    add_h2(doc, "Application Stack")
    headers = ["Layer", "Technology", "Why Chosen"]
    rows = [
        ["Runtime", "Node.js 22 LTS",
         "Long-term support, modern fetch and test runner, fast V8."],
        ["Language", "TypeScript 5.7 (strict)",
         "Compile-time guarantees on a sprawling route surface; zero use of any."],
        ["HTTP framework", "Fastify 5",
         "Two to three times the throughput of Express on equivalent "
         "hardware, schema-first plugins, first-class logging."],
        ["Templating", "EJS 3 (server-rendered)",
         "Plain HTML on the wire; no client bundle to parse before first paint."],
        ["Database", "SQLite via better-sqlite3 11",
         "Zero ops; file-backed; synchronous prepared statements; "
         "schema portable to Postgres."],
        ["Auth: passwords", "bcrypt (cost factor 12)",
         "Adaptive cost; widely audited; resistant to GPU brute force."],
        ["Auth: sessions", "JSON Web Tokens (HS256) in HttpOnly cookie",
         "Stateless verification with a server-side sessions table for revocation."],
        ["Validation", "zod",
         "Single source of truth for body, query, and param schemas; "
         "type-inferred handlers."],
        ["Logging", "Pino structured JSON",
         "One of the fastest Node loggers; native Cloud Logging shape."],
        ["Internationalisation",
         "Custom helper over en.json and hi.json",
         "No external dependency for two locales; cookie + Accept-Language "
         "resolution."],
        ["Images", "Sharp",
         "Native libvips bindings; sub-second WebP resize at 1600 x 1200."],
        ["Maps", "Leaflet + OpenStreetMap tiles",
         "No API key; lazy-loaded only on map pages; respects OSM tile "
         "usage policy."],
        ["IDs", "ULID",
         "Lexicographically sortable, collision-free, URL-safe."],
        ["CSRF",
         "@fastify/csrf-protection (double-submit cookie)",
         "Form-friendly; works with multipart; no shared session store needed."],
        ["Rate limiting", "@fastify/rate-limit",
         "200 requests per minute per IP by default; health endpoints "
         "allow-listed."],
        ["Container", "Multi-stage Dockerfile (Debian slim)",
         "Builder stage compiles native modules; runtime stage ships "
         "only libvips42 and libsqlite3-0."],
        ["Orchestration", "Google Cloud Run",
         "Scale to zero; per-request billing; native Cloud Logging "
         "integration; asia-south1 region for data residency."],
        ["Build", "Google Cloud Build",
         "One pipeline definition; same build runs locally and in CI."],
        ["Image registry", "Artifact Registry (asia-south1)",
         "Co-located with the Cloud Run service; vulnerability scanning."],
        ["Secrets", "Google Secret Manager",
         "Five secrets wired through cloudbuild.yaml; rotated independently."],
        ["Tests", "node:test + tests/smoke.mjs",
         "Zero extra dependencies; spawns the real server and exercises "
         "real HTTP routes."],
    ]
    add_table(doc, headers, rows,
              col_widths=[Cm(3.6), Cm(4.6), Cm(8.6)])
    add_caption(doc, "Table 4.1: Application stack and design rationale.")

    add_h2(doc, "Simulated External Systems and Production Replacements")
    headers2 = ["Simulated External System", "Real Production Replacement",
                "Reference"]
    rows2 = [
        ["DigiLocker OAuth 2.0 (PKCE) callback at "
         "/verify/digilocker/callback that issues a SIM_ code and "
         "marks kyc_verified = 1.",
         "Real DigiLocker partner credentials. Replace the SIM_ code "
         "exchange with a token POST to the DigiLocker token endpoint, "
         "fetch the e-Aadhaar payload, and store only sanitised fields. "
         "The full real-code block is recorded as a comment at the top of "
         "src/routes/verification.ts.",
         "https://digilocker.gov.in/"],
        ["Razorpay Orders API simulation in /pay/order that mints an "
         "order_<id> and returns it; webhook simulation in /pay/webhook "
         "that verifies an HMAC-SHA-256 signature and inserts a "
         "renter_records row on payment.captured.",
         "Real Razorpay account, real KEY_ID and KEY_SECRET, real "
         "WEBHOOK_SECRET. Replace the simulator with the official "
         "razorpay Node SDK as documented in the comment block at the "
         "top of src/routes/payments.ts.",
         "https://razorpay.com/docs/"],
        ["ONDC and Beckn protocol simulator under /ondc/v1/* covering "
         "search, select, init, confirm, status, cancel, update, rating, "
         "and support, with synchronous ACK or NACK and asynchronous "
         "on_* callbacks logged to ondc_messages.",
         "Register as an ONDC subscriber, generate Ed25519 and X25519 "
         "key pairs, configure the gateway URL, sign every outbound "
         "request per the Beckn auth header specification, and verify "
         "every inbound Authorization header against the registry.",
         "https://ondc.org/  |  https://developers.becknprotocol.io/"],
    ]
    add_table(doc, headers2, rows2,
              col_widths=[Cm(6.0), Cm(6.6), Cm(4.2)])
    add_caption(doc,
                "Table 4.2: Simulated external systems and the line edits "
                "required to switch each one to live operation.")

    add_h2(doc, "Architecture Diagram")
    add_para(doc,
             "[Figure 4.1: System architecture — Fastify single-container "
             "deployment on Cloud Run with ephemeral SQLite and local "
             "uploads, fronted by ONDC buyer applications and direct "
             "browser traffic.]",
             italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
             "[Figure 4.2: ONDC message flow — BAP search arrives at "
             "/ondc/v1/search, the BPP responds with synchronous ACK, then "
             "asynchronously POSTs on_search to context.bap_uri.]",
             italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4)
    add_para(doc,
             "Source SVG diagrams for both figures are tracked in the "
             "repository under /docs alongside the original draw.io files "
             "so reviewers can regenerate them.",
             italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11))


def build_methodology(doc):
    add_h1(doc, "Methodology")
    add_para(doc,
             "The capstone followed an iterative, single-author Minimum "
             "Lovable Product methodology. The unit of progress was a "
             "working end-to-end slice that could be demonstrated to a "
             "peer reviewer at the end of each week. There was no separate "
             "design phase, no thrown-away prototype, and no unfinished "
             "branch; each week extended a deployable application.")

    add_h2(doc, "Requirements Gathering")
    add_para(
        doc,
        "Formal user research was outside the scope of an academic "
        "capstone. In place of structured interviews, the requirements "
        "were modelled through three personas drawn from lived "
        "observation of campus housing in Solan and Mandi: Aarti, a "
        "first-year out-of-town student arriving with two suitcases and "
        "no broker contact; Rakesh, an owner of a six-room PG who "
        "currently relies on WhatsApp groups; and Priya, a hostel warden "
        "who needs to know which off-campus listings nearby are "
        "considered safe by her own students. Each persona's journey was "
        "sketched on paper, broken into screens, and then turned into a "
        "checklist of routes and database tables. The artefacts of this "
        "phase live in idea.md and SPEC.md.",
    )

    add_h2(doc, "Iterative Build with Working Software at Each Checkpoint")
    add_para(
        doc,
        "Each week ended with a working application deployed to a local "
        "port and demonstrable end-to-end. Week 3 ended with signup, "
        "login, and logout. Week 4 ended with a publishable listing. "
        "Week 6 ended with a complete student-to-owner booking. Week 8 "
        "ended with an end-to-end paid rental that automatically "
        "produced a renter_records row. Week 10 ended with a production "
        "Cloud Run URL. This rhythm meant that at every point in the "
        "twelve-week schedule there was a deployable artefact, not a "
        "half-finished branch.",
    )

    add_h2(doc, "Server-Rendered First, Then Enhanced")
    add_para(
        doc,
        "The application is server-rendered by default. Every page is "
        "fully usable with JavaScript disabled. Client-side enhancements "
        "are added only where they meaningfully improve the experience: "
        "the map on the search page is loaded lazily through a small "
        "intersection observer; the service worker caches the shell and "
        "the last twenty viewed listings; the locale switcher updates a "
        "cookie without a full reload. This order of work — "
        "server-side first, client-side strictly as enhancement — "
        "made it possible to ship a 2G-friendly application without "
        "negotiating with a frontend bundler.",
    )

    add_h2(doc, "Security and Accessibility as Cross-Cutting Concerns")
    add_para(
        doc,
        "Security and accessibility were treated as cross-cutting "
        "concerns enforced on every route, not as a final-week pass. "
        "Validation was added the moment a route was first written. CSRF "
        "tokens were wired before the first POST handler shipped. Every "
        "interactive element was tested with the keyboard before its "
        "screen was considered done. The result is that the WCAG 2.2 AAA "
        "audit in week 9 confirmed an existing baseline rather than "
        "fixing a backlog.",
    )

    add_h2(doc, "Testing Strategy")
    add_para(
        doc,
        "Tests were written as the surface they cover landed. The smoke "
        "suite spawns the real server in a child process against a "
        "temporary database and uploads directory, exercises real HTTP "
        "routes through native fetch, and tears the server down at the "
        "end. Every prod-shaped bug encountered during development became "
        "a regression test. The full suite runs in well under a minute "
        "and gates every push.",
    )


def build_work_done(doc):
    add_h1(doc, "Work Done")
    add_para(doc,
             "This section records the engineering work itself. I describe "
             "what I built, how it is wired, and why each non-obvious "
             "choice was made. File paths are given relative to the "
             "repository root.")

    add_h2(doc, "Architecture Overview")
    add_para(
        doc,
        "GharSetu runs as a single Node.js process inside a single "
        "container. Fastify 5 hosts the HTTP surface; EJS renders the "
        "HTML; better-sqlite3 holds the data; Pino emits structured logs "
        "to stdout, which Cloud Logging picks up. Static assets are "
        "served from /static; user-uploaded images from /uploads. ONDC "
        "buyer applications hit the same process at /ondc/v1/* and "
        "share the same database. There is no separate worker, no "
        "queue, and no cache. The text diagram in the README shows the "
        "full topology.",
    )

    add_h2(doc, "Data Model")
    add_para(doc,
             "The schema lives in src/db/schema.sql and is applied "
             "verbatim on cold start. Every table follows the same "
             "conventions: ULID primary key, epoch milliseconds for "
             "timestamps, integer rupees for money. The tables are:")
    add_bullets(doc, [
        "users — students, owners, and admins. Carries email, "
        "phone, password_hash, full_name, role, preferred_lang, KYC "
        "fields (kyc_verified, kyc_method, kyc_payload), and timestamps. "
        "Indexed on role.",
        "listings — the listing surface itself. Owner reference, "
        "title, description, property_type, gender_pref, rent_monthly, "
        "deposit, full address, lat / lng, near_landmark, amenities and "
        "rules as JSON arrays, available_from, status, view_count, "
        "timestamps. Indexed on (city, status), owner, and (lat, lng).",
        "listing_images — ordered images per listing, cascade-"
        "delete on parent removal.",
        "bookings — visit or reserve requests with status workflow "
        "(pending, accepted, declined, cancelled, completed) and an "
        "ondc_order_id slot for federation.",
        "payments — Razorpay shadow records: rzp_order_id (unique), "
        "rzp_payment_id, rzp_signature, status (created, captured, "
        "failed, refunded), and the for_month label.",
        "renter_records — the source of truth for the verified-"
        "renter flag. Source is either owner_marked or platform_payment. "
        "Unique on (listing_id, student_id).",
        "feedback — ratings and bodies, with the immutable "
        "is_verified_renter flag computed at insert time from the "
        "presence of an active renter_records row.",
        "audit_log — every mutation with actor, IP, user-agent, and "
        "a sanitised JSON payload. Indexed by created_at and by actor.",
        "ondc_messages — every Beckn message in either direction, "
        "indexed by (txn_id, created_at) for replay and debugging.",
        "sessions — JWT jti tracking for revocation, indexed on "
        "user_id.",
    ])

    add_h2(doc, "Authentication and Session Management")
    add_para(
        doc,
        "Passwords are hashed with bcrypt at cost factor 12. The "
        "password policy enforces a minimum length of ten characters and "
        "rejects entries from a small built-in list of common passwords. "
        "On successful login the server mints a JWT with claims "
        "{sub, role, jti, iat, exp}, signed with HS256 against the "
        "JWT_SECRET environment variable. The token is set as the "
        "gs_session cookie with HttpOnly, Secure, SameSite=Lax, and a "
        "seven-day max-age. The same jti is recorded in the sessions "
        "table so that a server-side revocation is possible without "
        "rotating the signing secret. CSRF protection uses the "
        "double-submit cookie pattern through @fastify/csrf-protection: "
        "the gs_csrf cookie is matched against a hidden _csrf form field "
        "on every state-changing POST.",
    )

    add_h2(doc, "Listing CRUD with Multipart Image Upload")
    add_para(
        doc,
        "Owners create and edit listings through HTML forms. Multipart "
        "uploads are handled by @fastify/multipart with a per-file limit "
        "of eight megabytes and a per-request limit of six files. Each "
        "incoming image is fed to Sharp, resized to a maximum bounding "
        "box of 1600 by 1200 pixels while preserving aspect, re-encoded "
        "to WebP at quality seventy-eight, and written to UPLOADS_DIR "
        "with a ULID filename. The resulting URL is recorded in "
        "listing_images with a position used for ordering on the detail "
        "page. Soft delete sets the listing status to removed; the row "
        "is preserved for audit and for the ondc_messages history.",
    )

    add_h2(doc, "Search and Discovery")
    add_para(
        doc,
        "The /search route accepts city, free-text query, rent range, "
        "property type, gender preference, amenities, sort order, and "
        "page parameters. Every parameter is validated through a zod "
        "schema before the database is touched. Queries are built with "
        "parameterised better-sqlite3 prepared statements; there is no "
        "string interpolation anywhere on the data path. Sort options "
        "include newest first, lowest rent first, and nearest first; the "
        "nearest-first sort uses the Haversine formula on the (lat, lng) "
        "column pair against the user-supplied geographic centre. The "
        "result set is rendered both as an HTML page (the default) and "
        "as a JSON response at /api/search. Map view is opt-in: Leaflet "
        "is loaded lazily only when the user toggles it.",
    )

    add_h2(doc, "Bookings and Owner Workflow")
    add_para(
        doc,
        "A student initiates a visit or reserve booking through the "
        "listing detail page. The POST /bookings handler verifies that "
        "the user is authenticated and that the listing is active before "
        "inserting a row with status pending. The booking shows up "
        "immediately in the owner's dashboard inbox. The owner accepts "
        "or declines through POST /bookings/:id/decision; the student "
        "sees the updated state in their own dashboard. An owner can "
        "subsequently mark a student as a renter on a listing through "
        "POST /owner/listings/:lid/renters; this inserts the "
        "owner_marked row in renter_records and unlocks the verified-"
        "renter feedback path.",
    )

    add_h2(doc, "Feedback System with Verified-Renter Logic")
    add_para(
        doc,
        "Feedback is open to any authenticated user. The "
        "is_verified_renter flag, however, is computed server-side at "
        "insert time from a single SQL probe: does an active row exist "
        "in renter_records with the matching listing_id and the author "
        "as student_id? If yes, the flag is set to 1. The author cannot "
        "change this flag through any HTML form or API parameter; the "
        "schema does not even expose it on the input zod object. "
        "Verified feedback is rendered with a distinct visual tag on "
        "the listing detail page and is sorted ahead of unverified "
        "feedback by default. The result is that a positive review by "
        "an outsider does not carry the same weight as a positive "
        "review by a former tenant, and an attempt to game the system "
        "by paying for one's own listing is, by construction, traceable "
        "in the payments table.",
    )

    add_h2(doc, "Razorpay Payment Simulation")
    add_para(
        doc,
        "POST /pay/order accepts a booking_id, listing_id, amount, and "
        "for_month. It mints an order identifier of the shape "
        "order_<ulid>, persists a payments row with status created, and "
        "returns { order_id, key_id, amount, currency } in the same "
        "JSON shape that a real Razorpay Orders API call would return. "
        "POST /pay/webhook receives the simulated webhook with the "
        "x-razorpay-signature header. The handler computes "
        "HMAC-SHA-256 over the raw request body using "
        "RZP_WEBHOOK_SECRET, compares it against the header in "
        "constant time, and rejects on mismatch. On the captured event "
        "the payments row moves to status captured, and the same "
        "transaction inserts an active row into renter_records with "
        "source set to platform_payment. The full real-shape replacement "
        "code is recorded in the comment block at the top of "
        "src/routes/payments.ts.",
    )

    add_h2(doc, "DigiLocker KYC Simulation")
    add_para(
        doc,
        "GET /verify renders the KYC start page for an authenticated "
        "user. POST /verify/digilocker/init mints a state value and a "
        "simulated authorization code prefixed SIM_, audits the kyc.init "
        "event, and redirects the browser to "
        "/verify/digilocker/callback. The callback validates the SIM_ "
        "prefix, builds a sanitised payload (method, verified_at, "
        "aadhaar_last4 = '1234', name_on_id, dob_year), and sets "
        "kyc_verified = 1, kyc_method = 'digilocker', and kyc_payload = "
        "<json> on the user row. No full Aadhaar number is ever stored. "
        "The full real OAuth 2.0 PKCE flow with the upstream DigiLocker "
        "token endpoint is recorded in the comment block at the top of "
        "src/routes/verification.ts.",
    )

    add_h2(doc, "ONDC Beckn Protocol Simulation")
    add_para(
        doc,
        "All nine endpoints live under /ondc/v1/. Every inbound payload "
        "is logged to ondc_messages with direction in. The handler "
        "responds synchronously with an ACK envelope of the shape "
        "{ message: { ack: { status: 'ACK' } } } if the payload "
        "validates, or with the corresponding NACK plus an error "
        "object if it does not. The asynchronous on_* reply is then "
        "dispatched in the background to context.bap_uri and logged "
        "with direction out. The action-to-domain mapping is direct: "
        "search returns matching listings as a Beckn catalog; select "
        "narrows to a single listing with a quote; init records "
        "billing intent on the booking; confirm transitions the "
        "booking to a firm reservation; status returns the current "
        "booking state; cancel and update mutate the booking with "
        "audit_log entries; rating bridges to the feedback table; "
        "support returns the platform's contact channels. The signing "
        "and registry-verification layer is documented as a single "
        "swap point in the source comments.",
    )

    add_h2(doc, "Internationalisation")
    add_para(
        doc,
        "Strings are stored in two JSON dictionaries at "
        "src/locales/en.json and src/locales/hi.json. The lib/render "
        "helper exposes a small t(key) function that the EJS templates "
        "call. Locale resolution per request follows a strict order: "
        "the ?lang= query parameter (which also rewrites the gs_lang "
        "cookie), the gs_lang cookie itself, the Accept-Language "
        "request header parsed for the first supported tag, and finally "
        "the DEFAULT_LANG environment variable. The html lang attribute "
        "is set per response so screen readers pronounce content "
        "correctly. The dir attribute is set per locale, so adding a "
        "right-to-left language such as Urdu is a JSON-only change with "
        "no layout work.",
    )

    add_h2(doc, "Accessibility (WCAG 2.2 AAA)")
    add_para(
        doc,
        "Every text colour pair on the site was tested with a contrast "
        "checker against the AAA bar of 7:1; the brand saffron was "
        "darkened until both light and dark backgrounds passed. The "
        "first focusable element on every page is a Skip to content "
        "link that is hidden until it receives focus. Focus rings are "
        "visible through :focus-visible at a contrast ratio of at least "
        "3:1 against any background. Icon-only buttons carry an "
        "aria-label. Form errors are announced to assistive technology "
        "through aria-describedby pointing at a sibling error span. "
        "The prefers-reduced-motion: reduce media query disables every "
        "transition. The forced-colors media query is respected so that "
        "Windows High Contrast users see system colours. Live testing "
        "was carried out with NVDA on Firefox and VoiceOver on Safari.",
    )

    add_h2(doc, "Observability")
    add_para(
        doc,
        "Pino emits one structured JSON line per request to stdout, "
        "including the ULID req.id, method, path, status, response time "
        "in milliseconds, the resolved user_id (if any), the client IP, "
        "and the user-agent. Cloud Logging picks these up automatically "
        "in production. Every state-changing request also writes to the "
        "audit_log table through a centralised audit() helper invoked "
        "from the onResponse hook in src/server.ts. The /admin route is "
        "an SIEM-style page that streams the latest audit entries with "
        "a filterable timestamp / actor / action UI. Health is exposed "
        "at /healthz (always 200 if the process is up) and at /readyz "
        "(checks database connectivity). Both are excluded from rate "
        "limiting.",
    )

    add_h2(doc, "Containerisation and Cloud Run Deployment")
    add_para(
        doc,
        "The Dockerfile is multi-stage. The builder stage runs on "
        "node:22-bookworm-slim with python3, make, g++, pkg-config, "
        "libvips-dev, and libsqlite3-dev installed for the native "
        "module compile. It installs all dependencies, transpiles the "
        "TypeScript, prunes dev dependencies, and clears the npm "
        "cache. The runtime stage starts from the same slim image with "
        "only libvips42, libsqlite3-0, ca-certificates, tini, and curl "
        "installed; it copies the pruned node_modules and the built "
        "dist directory, runs as the unprivileged node user, exposes "
        "port 8080, and ships a HEALTHCHECK directive against /healthz. "
        "The cloudbuild.yaml file builds, pushes to Artifact Registry "
        "with both :$SHORT_SHA and :latest tags, and deploys to Cloud "
        "Run with min-instances 0, max-instances 10, 512 MiB of memory, "
        "1 vCPU, concurrency 80, and the five secrets wired through "
        "--set-secrets. The deploy.sh script is idempotent: it enables "
        "every required Google Cloud API, creates the Artifact Registry "
        "repository, creates and IAM-binds the five secrets if they do "
        "not exist, grants Cloud Build the necessary roles on the "
        "runtime service account, and submits the build.",
    )

    add_h2(doc, "Testing")
    add_para(doc,
             "The smoke suite at tests/smoke.mjs spawns the real server in "
             "a child process with a temporary SQLite file and uploads "
             "directory and exercises the live HTTP surface. It contains "
             "the following ten checks, each of which gates every push:")
    add_numbered(doc, [
        "GET /healthz returns 200 with { ok: true }.",
        "GET / returns HTML containing the brand string GharSetu.",
        "GET /search returns HTML containing at least one Shoolini-area "
        "listing from the seeded dataset.",
        "GET /api/search?city=Solan returns a JSON object whose data "
        "array is non-empty.",
        "Signup followed by login issues a gs_session cookie under a "
        "fresh user account, end to end, with CSRF tokens.",
        "An authenticated student can POST /bookings against a known "
        "listing and receive a successful response.",
        "POST /pay/order returns an order identifier prefixed with "
        "order_, matching the real Razorpay Orders API shape.",
        "POST /ondc/v1/search with a valid Beckn search envelope "
        "returns 200 with message.ack.status equal to ACK.",
        "GET /admin without an admin cookie is blocked with 401, 403, "
        "or a redirect.",
        "GET /healthz returns the documented body shape, re-checked "
        "after the suite has exercised every other route.",
    ])
    add_para(doc,
             "All ten tests pass on every run. Every prod-shaped bug "
             "encountered during development was added here as a "
             "regression test before the fix landed.")


def build_challenges(doc):
    add_h1(doc, "Challenges and Solutions")

    add_h2(doc, "Fastify v5 Logger Type Incompatibility with Pino")
    add_para(
        doc,
        "Fastify 5 expects a logger conforming to FastifyBaseLogger, "
        "while Pino's exported Logger type is a near-miss superset. The "
        "TypeScript compile failed at registration. Resolved by passing "
        "the Pino instance through an explicit cast at the single "
        "registration site in src/server.ts, with a comment recording "
        "the version pair (Fastify 5.x and Pino 9.x) at which the cast "
        "is required. No runtime surface was changed.",
    )

    add_h2(doc, "Ephemeral SQLite on Cloud Run /tmp")
    add_para(
        doc,
        "Cloud Run's /tmp is in-memory tmpfs, scoped to a single "
        "instance, and lost on scale-to-zero. For a public capstone "
        "demo where users might create a listing and return the next "
        "day, this is an obvious footgun. Resolved by reseeding on "
        "every cold start (so the demo data is always present) and "
        "documenting the constraint prominently in the README. The "
        "production migration path — Cloud SQL for PostgreSQL with "
        "the same schema and the better-sqlite3 driver swapped for "
        "pg — is recorded in the same place so the swap is a "
        "credentials and connection-string change rather than a "
        "schema rewrite.",
    )

    add_h2(doc, "WCAG AAA Contrast Versus Brand Colour")
    add_para(
        doc,
        "The first iteration of the brand palette used a saffron "
        "(#FF9933) that failed the AAA bar of 7:1 against white "
        "backgrounds. Resolved by darkening the saffron to a value that "
        "passes 7:1 on white and 4.5:1 on dark, and by reserving the "
        "lighter saffron for non-text accents only. Verified with a "
        "low-vision simulation filter and a side-by-side check against "
        "a colour-blindness simulator.",
    )

    add_h2(doc, "Beckn Asynchronous Reply Reliability")
    add_para(
        doc,
        "The Beckn protocol expects a synchronous ACK followed by an "
        "asynchronous on_* callback to context.bap_uri. A naive "
        "implementation that does the reply in-line risks holding the "
        "ACK request open across a slow upstream call. Resolved by "
        "logging the inbound message, returning the ACK immediately, "
        "and dispatching the on_* reply through a fire-and-forget "
        "promise that catches and logs every error. Both messages are "
        "persisted to ondc_messages, so a missed callback can be "
        "replayed by hand from the database.",
    )

    add_h2(doc, "Image Weight on Slow Networks")
    add_para(
        doc,
        "Owner-uploaded photos arrive in arbitrary sizes, often above "
        "five megabytes. Serving them as-is would destroy the 2G budget. "
        "Resolved at upload time: every image is fed to Sharp, resized "
        "to a maximum bounding box of 1600 by 1200 pixels while "
        "preserving aspect, and re-encoded to WebP at quality "
        "seventy-eight. The cap of six images per listing keeps the "
        "total page weight under one megabyte even on the heaviest "
        "listing in the seed.",
    )

    add_h2(doc, "CSRF Plus Multipart Form Interaction")
    add_para(
        doc,
        "@fastify/csrf-protection's default behaviour expects a JSON or "
        "URL-encoded body, but listing creation has to be multipart for "
        "image upload. The naive integration rejected every multipart "
        "POST. Resolved by reading the _csrf field through the multipart "
        "parser (where it is exposed as a non-file field), validating "
        "it against the gs_csrf cookie before the file fields are "
        "processed, and rejecting with a friendly error page if the "
        "match fails.",
    )

    add_h2(doc, "Native Module Compile Time in the Container Build")
    add_para(
        doc,
        "Both better-sqlite3 and Sharp ship native bindings that take "
        "well over a minute to compile inside Docker. A single-stage "
        "build dragged build times above five minutes and bloated the "
        "runtime image with toolchains. Resolved by adopting a "
        "multi-stage Dockerfile: the builder stage installs the "
        "build-essential tooling, compiles the natives, and runs npm "
        "prune --omit=dev; the runtime stage installs only libvips42 "
        "and libsqlite3-0 and copies the pruned node_modules across. "
        "The final image weight dropped by more than half and the cold "
        "start improved correspondingly.",
    )


def build_learning_outcomes(doc):
    add_h1(doc, "Learning Outcomes")
    add_para(
        doc,
        "The capstone consolidated several skill areas into a single "
        "deployable artefact. The most concrete outcomes are listed "
        "below.",
    )
    add_bullets(doc, [
        "End-to-end TypeScript on the server, with strict mode enabled "
        "and zero use of any across the codebase, including in the "
        "Beckn envelope handlers.",
        "Cloud-native deployment as a single container on Google Cloud "
        "Run, with cold starts under five seconds, idle billing at zero "
        "rupees, and a production-ready CI/CD pipeline through Cloud "
        "Build and Artifact Registry.",
        "Working knowledge of the ONDC and Beckn protocol at the "
        "message-flow level, including the synchronous-ACK + "
        "asynchronous-callback pattern, the registry and signing model, "
        "and the retail-services profile actions.",
        "Accessibility engineering against WCAG 2.2 AAA contrast on "
        "every text element, with live verification on NVDA and "
        "VoiceOver and a complete keyboard-only journey from landing "
        "page to confirmed booking.",
        "A security-by-default mindset: bcrypt at cost twelve, JWTs "
        "with revocation through a sessions table, double-submit CSRF, "
        "rate limiting, parameterised SQL on every data path, and a "
        "database-backed audit log on every mutation.",
        "Pragmatic Docker authorship through a multi-stage build that "
        "compiles native modules in a builder stage and ships only the "
        "necessary shared libraries in the runtime stage.",
        "Practical internationalisation through a small custom resolver "
        "that respects an explicit query, a stored cookie, the request "
        "header, and a default in that order, with right-to-left "
        "readiness through the dir attribute.",
        "Technical writing discipline: a SPEC.md that committed the "
        "scope before the first line of code was written, a README.md "
        "that records the production swap path next to every "
        "simulation, and this report.",
        "Team-style coordination as a single author by treating each "
        "weekly slice as if it were a pull request to be reviewed, "
        "merged, and demonstrated.",
    ])


def build_conclusion(doc):
    add_h1(doc, "Conclusion")
    add_para(
        doc,
        "GharSetu set out to demonstrate that an open, federated, "
        "student-first housing platform can be built and deployed by a "
        "single final-year student on infrastructure that costs nothing "
        "at idle. The result is a working application with a verified-"
        "renter trust mechanism, a faithful Beckn protocol simulator "
        "covering nine actions, real-shape DigiLocker and Razorpay "
        "integrations, an admin SIEM dashboard, English and Hindi "
        "localisation, WCAG 2.2 AAA contrast on every text element, ten "
        "smoke tests passing on every run, and a Cloud Run deployment "
        "whose cold start completes in under five seconds.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "What this project does not yet have is also worth recording "
        "honestly. It is not a registered ONDC subscriber; the Beckn "
        "endpoints are simulated in-process, and a real subscription "
        "with Ed25519 key generation, registry lookup, and request "
        "signing is the obvious next step. The database is ephemeral on "
        "Cloud Run /tmp; a Cloud SQL for PostgreSQL migration is the "
        "next obvious step there, and the schema is already portable. "
        "There is no native mobile application yet, although the web "
        "application is installable as a progressive web app through "
        "the included service worker and manifest. Longer term, the "
        "verified-renter model could be extended with neighbourhood-"
        "level verification (a hostel warden vouching for a property, "
        "for example), and the federation could carry not just listings "
        "but also dispute resolution.",
        first_line_indent=Cm(1.0),
    )
    add_para(
        doc,
        "If the project demonstrates anything, it is that the building "
        "blocks for an honest, federated, locally-anchored Indian "
        "student housing market are now within reach of a single "
        "developer. Open standards, ephemeral compute, and a small set "
        "of well-maintained open-source libraries do most of the work. "
        "What remains is the patient business of earning trust, owner "
        "by owner and student by student, in the small towns where the "
        "supply already lives.",
        first_line_indent=Cm(1.0),
    )


def build_references(doc):
    add_h1(doc, "References")
    refs = [
        "Open Network for Digital Commerce. (2024). About ONDC. "
        "Retrieved from https://ondc.org/",
        "Beckn Foundation. (2024). Beckn Protocol Specifications and "
        "Developer Documentation. Retrieved from "
        "https://developers.becknprotocol.io/",
        "Government of India, Ministry of Electronics and Information "
        "Technology. (2024). DigiLocker Developer Documentation. "
        "Retrieved from https://digilocker.gov.in/",
        "Razorpay Software Pvt. Ltd. (2024). Razorpay Developer "
        "Documentation — Orders API and Webhooks. Retrieved from "
        "https://razorpay.com/docs/",
        "World Wide Web Consortium. (2023). Web Content Accessibility "
        "Guidelines (WCAG) 2.2 — Quick Reference. Retrieved from "
        "https://www.w3.org/WAI/WCAG22/quickref/",
        "Google Cloud. (2024). Cloud Run Documentation. Retrieved from "
        "https://cloud.google.com/run/docs",
        "Google Cloud. (2024). Cloud Build Documentation. Retrieved "
        "from https://cloud.google.com/build/docs",
        "Google Cloud. (2024). Secret Manager Documentation. Retrieved "
        "from https://cloud.google.com/secret-manager/docs",
        "Fastify. (2024). Fastify Documentation. Retrieved from "
        "https://fastify.dev/docs/latest/",
        "Pino. (2024). Pino Logger Documentation. Retrieved from "
        "https://getpino.io/",
        "Mozilla Developer Network. (2024). Service Worker API. "
        "Retrieved from "
        "https://developer.mozilla.org/en-US/docs/Web/API/Service_Worker_API",
        "OpenStreetMap Foundation. (2024). Tile Usage Policy. "
        "Retrieved from "
        "https://operations.osmfoundation.org/policies/tiles/",
        "Government of India, Ministry of Electronics and Information "
        "Technology. (2023). The Digital Personal Data Protection Act, "
        "2023. Retrieved from https://meity.gov.in/",
        "Government of India, Ministry of Education. (2024). All India "
        "Survey on Higher Education (AISHE) Report. Retrieved from "
        "https://aishe.gov.in/",
        "JLL India. (2024). Student Housing Market Report — India. "
        "Retrieved from https://www.jll.co.in/",
        "Volkov, R. (2024). better-sqlite3 — The fastest and "
        "simplest library for SQLite3 in Node.js. Retrieved from "
        "https://github.com/WiseLibs/better-sqlite3",
        "Lovell Fuller. (2024). Sharp — High performance Node.js "
        "image processing. Retrieved from https://sharp.pixelplumbing.com/",
    ]
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        _set_para_spacing(p, line=1.3, before=0, after=4,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        run = p.add_run(f"[{i}]  ")
        _set_run_font(run, size=BODY_SIZE, bold=True)
        run2 = p.add_run(ref)
        _set_run_font(run2, size=BODY_SIZE)


# ---------- main -------------------------------------------------------------


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)

    build_cover_page(doc)
    build_declaration(doc)
    build_certificate(doc)
    build_project_diary(doc)
    build_supervisor_feedback(doc)
    build_table_of_contents(doc)
    build_acknowledgement(doc)
    build_project_profile(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_objectives(doc)
    build_theoretical_background(doc)
    build_tools(doc)
    build_methodology(doc)
    build_work_done(doc)
    build_challenges(doc)
    build_learning_outcomes(doc)
    build_conclusion(doc)
    build_references(doc)

    return doc


def report_section_structure(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        sty = p.style.name if p.style else ""
        if sty in {"Heading 1", "Heading 2"}:
            prefix = "#" if sty == "Heading 1" else "  ##"
            yield f"{prefix} {p.text}"


def main() -> int:
    doc = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))

    size_bytes = OUTPUT_PATH.stat().st_size
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    print("=" * 70)
    print(f"FILE      : {OUTPUT_PATH}")
    print(f"SIZE      : {size_bytes:,} bytes ({size_bytes / 1024:.1f} KiB)")
    print(f"PARAGRAPHS: {para_count}")
    print(f"TABLES    : {table_count}")
    print("=" * 70)
    print("STRUCTURE :")
    for line in report_section_structure(doc):
        print(line)
    print("=" * 70)
    return 0


if __name__ == "__main__":
    sys.exit(main())
